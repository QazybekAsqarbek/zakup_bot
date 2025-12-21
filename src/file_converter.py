import io
import logging
import pandas as pd
import docx

logger = logging.getLogger(__name__)

def convert_file_to_text(file_bytes: bytes, filename: str) -> str:
    """
    Принимает байты файла и имя файла.
    Возвращает текстовое представление содержимого.
    """
    filename = filename.lower()
    
    try:
        # 1. Обработка Word (.docx)
        if filename.endswith('.docx'):
            logger.info(f"🔄 Converting DOCX: {filename}")
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Также вытаскиваем таблицы из Word, это важно для КП!
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)

        # 2. Обработка Excel (.xlsx, .xls)
        elif filename.endswith(('.xlsx', '.xls')):
            logger.info(f"🔄 Converting EXCEL: {filename}")
            # Читаем Excel в DataFrame
            df = pd.read_excel(io.BytesIO(file_bytes))
            # Конвертируем в Markdown таблицу (DeepSeek её отлично понимает)
            return df.to_markdown(index=False)

        # 3. Обработка CSV
        elif filename.endswith('.csv'):
            logger.info(f"🔄 Converting CSV: {filename}")
            df = pd.read_csv(io.BytesIO(file_bytes))
            return df.to_markdown(index=False)

        # 4. Текстовые файлы (.txt, .md)
        elif filename.endswith(('.txt', '.md', '.py', '.json')):
            logger.info(f"🔄 Reading Text file: {filename}")
            return file_bytes.decode('utf-8')

        else:
            logger.warning(f"⚠️ Unknown file extension: {filename}")
            return None

    except Exception as e:
        logger.error(f"❌ File conversion error for {filename}: {e}", exc_info=True)
        return None
